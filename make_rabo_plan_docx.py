from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = "Velo_Vital_Rabo_Coöperatiefonds_Plan_van_Aanpak_en_Begroting.docx"

BLUE = RGBColor(33, 74, 130)
DEEP = RGBColor(38, 41, 64)
MUTED = RGBColor(95, 96, 110)
ORANGE = RGBColor(226, 105, 47)
LIGHT = "F7F4F6"
LIGHT_BLUE = "EEF3F8"
GRID = "D9DCE3"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_tbl_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def set_column_widths(table, widths_dxa):
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def paragraph_border_bottom(paragraph, color="E2692F", size="10", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_p(doc, text="", style=None, after=6, before=0, line=1.25, color=DEEP, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if text:
        r = p.add_run(text)
        set_run(r, size=11, color=color, bold=bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run(r, size=10.7, color=DEEP)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run(run, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run(run, size=12, color=DEEP, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_tbl_width(table, 9360, 120)
    set_table_borders(table, "E6DDE5", "4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, 150, 190, 150, 190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + " ")
    set_run(r1, size=10.5, color=ORANGE, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=10.5, color=DEEP)
    add_p(doc, "", after=4)


def add_info_table(doc):
    rows = [
        ("Organisatie", "Stichting Wielerevenementen Aalburg"),
        ("Project", "Velo Vital Leenfiets & Community Programma"),
        ("Rechtsvorm", "Stichting"),
        ("KvK-nummer", "60618167"),
        ("Projectgebied", "'s-Hertogenbosch en omstreken"),
        ("Projectperiode", "September 2026 tot en met augustus 2027"),
        ("Gevraagde bijdrage", "EUR 15.000"),
        ("Totale projectwaarde", "EUR 30.000"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_tbl_width(table, 9360, 120)
    set_column_widths(table, [2500, 6860])
    set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        set_cell_shading(c1, LIGHT_BLUE)
        set_cell_margins(c1)
        set_cell_margins(c2)
        for c in (c1, c2):
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run(r1, size=10.3, color=BLUE, bold=True)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run(r2, size=10.3, color=DEEP, bold=(label in ("Gevraagde bijdrage", "Totale projectwaarde")))
    add_p(doc, "", after=8)


def add_schedule_table(doc):
    rows = [
        ("Juli - augustus 2026", "Voorbereiding, partnerafspraken, selectie fietsen, opzetten aanmeldproces, campagnevoorbereiding."),
        ("September - november 2026", "Start lokale campagne, werving deelnemers, intakegesprekken, eerste leenfietsronde en community-activiteiten."),
        ("December 2026 - februari 2027", "Tweede leenfietsronde, winterproof community-aanbod, begeleiding, onderhoud en tussentijdse evaluatie."),
        ("Maart - mei 2027", "Derde leenfietsronde, extra kennismakingsmomenten, doorstroom naar structureel aanbod en partneractivatie."),
        ("Juni - augustus 2027", "Afronding jaarpilot, impactmeting, financiële afronding, terugkoppeling aan financiers en uitwerking van opschalingslessen."),
    ]
    add_table(doc, ["Periode", "Activiteiten"], rows, [2600, 6760])


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_tbl_width(table, 9360, 120)
    set_column_widths(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run(r, size=10.2, color=BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_shading(cells[idx], WHITE)
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            is_money = idx == len(row) - 1 and str(value).startswith("EUR")
            set_run(r, size=9.8, color=DEEP, bold=is_money)
            if is_money:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_p(doc, "", after=6)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Calibri"

    header = section.header.paragraphs[0]
    header.text = "Velo Vital | Rabo Coöperatiefonds"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.text = "Plan van aanpak en begroting"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    add_p(doc, "Stichting Wielerevenementen Aalburg", after=8, color=MUTED, bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Velo Vital Leenfiets & Community Programma")
    set_run(r, size=24, color=DEEP, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Plan van aanpak en begroting voor het Rabo Coöperatiefonds")
    set_run(r, size=13, color=MUTED, bold=True)
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, "E2692F", "12", "8")
    add_p(doc, "", after=4)
    add_info_table(doc)
    add_callout(
        doc,
        "Kern:",
        "Velo Vital helpt vrouwen in 's-Hertogenbosch en omgeving laagdrempelig kennismaken met sportief fietsen, met leenfietsen, begeleiding en een toegankelijke community als opstap naar structureel bewegen.",
    )

    add_heading(doc, "1. Doel van de organisatie")
    add_p(doc, "Stichting Wielerevenementen Aalburg zet zich met Velo Vital in om inwoners laagdrempelig in beweging te brengen via sportief fietsen, community-activiteiten en sociale verbinding. De organisatie wil drempels verlagen voor mensen die wel willen bewegen, maar nog niet vanzelf aansluiting vinden bij bestaand sportaanbod.")
    add_p(doc, "Met Velo Vital bouwen we aan een toegankelijke fietscommunity waarin gezondheid, ontmoeting, zelfvertrouwen en duurzame mobiliteit centraal staan. De stichting gebruikt haar ervaring met wielerevenementen en community-opbouw om niet alleen eenmalige activiteiten te organiseren, maar juist structurele deelname aan bewegen te stimuleren.")

    add_heading(doc, "2. Historie en doorontwikkeling")
    add_p(doc, "Velo Vital komt voort uit meer dan 10 jaar ervaring met het Marianne Vos Wielerfestival. Marianne Vos was daarin het herkenbare boegbeeld van vrouwenfietsen: sportief, toegankelijk en inspirerend voor recreanten, talenten en liefhebbers van de wielersport.")
    add_p(doc, "De stichting heeft in die periode veel ervaring opgebouwd met vrouwenfietssport, recreatieve deelname en aanwas richting de sport. Er zijn meerdere wielertoertochten en fietsevenementen georganiseerd, met in totaal meer dan 7.500 deelnemers. Daarmee is een stevig netwerk ontstaan rond vrouwen, fietsen, vrijwilligers, partners en lokale zichtbaarheid.")
    add_p(doc, "Tegelijk zien we dat de markt en de behoefte van deelnemers veranderen. Waar eerder een evenement of toertocht een logische instap was, vraagt de huidige doelgroep vaker om een laagdrempelige, persoonlijke en flexibele route naar structureel bewegen. Veel vrouwen willen eerst proberen, materiaal ervaren, vertrouwen opbouwen en aansluiting vinden bij een community voordat zij zich verbinden aan een club, evenement of eigen investering.")
    add_p(doc, "Daarom ontwikkelt de stichting haar ervaring door naar Velo Vital. Met de beschikbare middelen investeren we niet in een eenmalig evenement, maar in een nieuw maatschappelijk initiatief dat beter past bij de huidige behoefte: een leenfiets- en communityprogramma dat vrouwen helpt om duurzaam in beweging te komen.")
    add_p(doc, "De eerste focus ligt bewust op regio 's-Hertogenbosch. Deze lokale uitvoering is bedoeld als pilot: we testen en verbeteren hier het model voor werving, intake, leenfietsen, begeleiding, communityvorming, onderhoud en doorstroom. Als de aanpak werkt, wil Velo Vital het initiatief de komende jaren ook in andere regio's in Nederland uitrollen.")

    add_heading(doc, "3. Aanleiding")
    add_p(doc, "Sportief fietsen combineert beweging, buiten zijn, gezondheid, vrijheid en sociale ontmoeting. Toch is de stap om te beginnen voor veel vrouwen groot. Anders dan bij veel andere sporten vraagt fietsen vaak direct om materiaal, kennis, vertrouwen en een netwerk.")
    add_bullets(doc, [
        "hoge opstartkosten voor fiets, helm, kleding en accessoires;",
        "onzekerheid over materiaal en onderhoud;",
        "geen fietsmaatje of groep om mee te starten;",
        "twijfel of sportief fietsen past bij conditie, ervaring of leefritme;",
        "behoefte aan een toegankelijke, niet-prestatiegerichte instap;",
        "een sportcultuur die voor beginners soms technisch of minder toegankelijk voelt.",
    ])
    add_p(doc, "Velo Vital speelt hier direct op in: deelnemers hoeven niet eerst te kopen, te presteren of lid te worden van een club. Ze kunnen op een laagdrempelige manier proberen, ervaren en daarna zelfverzekerd verder.")

    add_heading(doc, "4. Doel van het project")
    add_p(doc, "Het doel van het Velo Vital Leenfiets & Community Programma is om vrouwen in 's-Hertogenbosch en omgeving laagdrempelig te laten kennismaken met sportief fietsen en hen te begeleiden naar structureel beweeggedrag.")
    add_bullets(doc, [
        "gezond leven en meer bewegen;",
        "sociale verbinding en ontmoeting;",
        "sterke lokale communities;",
        "duurzame mobiliteit;",
        "het verlagen van praktische, sociale en financiële drempels;",
        "doorstroom naar bestaand lokaal sport- en beweegaanbod.",
    ])

    add_heading(doc, "5. Doelgroep")
    add_p(doc, "De primaire doelgroep bestaat uit volwassen vrouwen in 's-Hertogenbosch en omgeving die interesse hebben in sportief fietsen, maar de stap nog niet zetten.")
    add_bullets(doc, [
        "weinig of onregelmatig sporten;",
        "willen werken aan conditie, vitaliteit of zelfvertrouwen;",
        "geen eigen geschikte fiets hebben;",
        "twijfelen of sportief fietsen bij hen past;",
        "graag samen met anderen willen bewegen;",
        "behoefte hebben aan een veilige, toegankelijke en niet-prestatiegerichte instap.",
    ])
    add_p(doc, "Binnen de communicatie kan ook expliciet ruimte worden gemaakt voor jonge vrouwen vanaf 16 jaar die dreigen af te haken bij sport of moeite hebben om passend sportaanbod te vinden.")

    add_heading(doc, "6. Projectaanpak")
    approach = [
        ("Leenfietsprogramma", "Velo Vital stelt fietsen beschikbaar aan vrouwen die sportief fietsen willen proberen. Deelnemers kunnen een fiets tijdelijk gebruiken voor een periode van een tot drie maanden."),
        ("Intake en matching", "Iedere deelnemer krijgt een korte intake. Daarbij kijken we naar ervaring, motivatie, beschikbare tijd, eventuele onzekerheden, gewenste intensiteit en praktische behoeften."),
        ("Community en fietsmomenten", "Deelnemers worden onderdeel van de Velo Vital-community. Zij kunnen aansluiten bij laagdrempelige fietsmomenten en contact leggen met andere vrouwen."),
        ("Fietschecks, onderhoud en materiaal", "Velo Vital werkt met een fietsenmaker of materiaalpartner die fietsen beschikbaar stelt, onderhoud uitvoert en zorgt dat het materiaal veilig en gebruiksklaar blijft."),
        ("Lokale campagne en zichtbaarheid", "Er wordt een zichtbaarheidscampagne opgezet om vrouwen in 's-Hertogenbosch en omgeving te bereiken die sportief fietsen interessant vinden, maar nog niet instappen."),
        ("Doorstroom naar blijvend bewegen", "Aan het einde van de leenperiode bespreekt Velo Vital met iedere deelnemer wat een passende vervolgstap is: eigen fiets, Velo Vital, Vrouwenwielrennen Den Bosch of zelfstandig verder fietsen met een fietsmaatje."),
    ]
    for label, text in approach:
        add_heading(doc, label, 3)
        add_p(doc, text)

    add_heading(doc, "7. Beoogde resultaten")
    add_p(doc, "In de projectperiode van een jaar wil Velo Vital de volgende resultaten bereiken:")
    add_bullets(doc, [
        "minimaal 25 vrouwen nemen deel aan een leenfietsperiode van een tot drie maanden;",
        "minimaal 75 vrouwen worden actief betrokken via community-activiteiten, kennismakingsmomenten, proefritten of campagne;",
        "minimaal 500 vrouwen worden bereikt via lokale communicatie, sociale media, partners en mond-tot-mondwerving;",
        "minimaal 20 van de 25 leenfietsdeelnemers blijven na afloop actief fietsen;",
        "minimaal 12 deelnemers stromen door naar structureel aanbod;",
        "minimaal 12 laagdrempelige community- of kennismakingsmomenten worden georganiseerd in regio 's-Hertogenbosch;",
        "er wordt een herhaalbare pilotaanpak ontwikkeld voor leenfietsen, intake, communitybinding, onderhoud, doorstroom en latere opschaling naar andere regio's.",
    ])

    add_heading(doc, "8. Planning")
    add_schedule_table(doc)

    add_heading(doc, "9. Samenwerking en partners")
    add_p(doc, "Velo Vital werkt met lokale en regionale partners om deelnemers te bereiken, kwaliteit te borgen en doorstroom mogelijk te maken.")
    add_bullets(doc, [
        "fietsenmaker of materiaalpartner voor fietsen, onderhoud en technische expertise;",
        "Vrouwenwielrennen Den Bosch als mogelijke doorstroompartner;",
        "S-PORT en Bosch Sportakkoord voor lokale sportverbinding en zichtbaarheid;",
        "Den Bosch City voor lokale zichtbaarheid, netwerkbereik en promotionele ondersteuning;",
        "wijk-, welzijns- en preventiepartners voor het bereiken van vrouwen die niet vanzelf bij sportaanbod terechtkomen;",
        "vrijwilligers en communitybegeleiders voor laagdrempelige begeleiding en fietsmomenten.",
    ])

    add_heading(doc, "10. Publiciteit")
    add_p(doc, "Het project wordt actief in de publiciteit gebracht via sociale media, deelnemersverhalen, lokale partnerkanalen, community-activiteiten, lokale netwerken en mond-tot-mondwerving.")
    add_p(doc, "Bij een bijdrage van Rabobank benoemt Velo Vital Rabobank als maatschappelijke partner in de communicatie. Velo Vital deelt voortgang, beeldmateriaal en resultaten met Rabobank en draagt actief bij aan positieve publiciteit rondom de bijdrage.")

    add_heading(doc, "11. Monitoring en evaluatie")
    add_p(doc, "Velo Vital meet de voortgang op een praktische manier. We registreren:")
    add_bullets(doc, [
        "aantal aanmeldingen;",
        "aantal deelnemers dat daadwerkelijk een fiets leent;",
        "duur van de leenperiode per deelnemer;",
        "deelname aan community-activiteiten;",
        "aantal deelnemers dat na afloop blijft fietsen;",
        "aantal deelnemers dat doorstroomt naar Velo Vital, Vrouwenwielrennen Den Bosch of ander aanbod;",
        "ervaringen van deelnemers via korte testimonials;",
        "leerpunten rond materiaal, begeleiding, onderhoud en werving.",
    ])

    add_heading(doc, "12. Begroting")
    add_heading(doc, "Kosten", 2)
    cost_rows = [
        ("Lokale marketingcampagne en zichtbaarheid voor deelnemerswerving", "EUR 5.000"),
        ("Content, vormgeving en deelnemersverhalen voor social media en partners", "EUR 3.000"),
        ("Kennismakingsritten en community-activatie voor nieuwe deelnemers", "EUR 3.500"),
        ("Fietschecks, onderhoud en rijklaar maken van leenfietsen", "EUR 4.500"),
        ("Praktisch instapmateriaal voor deelnemers", "EUR 3.000"),
        ("Aanmeldproces, intake en praktische deelnemersondersteuning", "EUR 2.500"),
        ("Monitoring, evaluatie en korte impactrapportage", "EUR 1.000"),
        ("Inzet fietsenmaker of materiaalpartner in natura", "EUR 3.500"),
        ("Vrijwilligersinzet community en begeleiding in natura", "EUR 2.000"),
        ("Inzet lokale partners voor zichtbaarheid, bereik en promotie in natura", "EUR 1.500"),
        ("Eigen inzet Velo Vital: organisatie, afstemming en communitybeheer", "EUR 500"),
        ("Totale projectwaarde", "EUR 30.000"),
    ]
    add_table(doc, ["Kostenpost", "Bedrag"], cost_rows, [7360, 2000])
    add_heading(doc, "Dekking", 2)
    income_rows = [
        ("Gevraagde bijdrage Rabo Coöperatiefonds", "EUR 15.000"),
        ("Eigen inzet Velo Vital: organisatie, afstemming en communitybeheer", "EUR 500"),
        ("Bijdrage fietsenmaker of materiaalpartner in natura", "EUR 3.500"),
        ("Vrijwilligersinzet community en begeleiding", "EUR 2.000"),
        ("Bijdrage lokale partners in natura: bereik, zichtbaarheid en promotie", "EUR 1.500"),
        ("Aangevraagde of beoogde bijdrage sport- en maatschappelijke fondsen", "EUR 7.500"),
        ("Totale dekking", "EUR 30.000"),
    ]
    add_table(doc, ["Dekking", "Bedrag"], income_rows, [7360, 2000])
    add_callout(
        doc,
        "Toelichting:",
        "Rabobank is niet de enige financier. Het project wordt mede mogelijk gemaakt door eigen inzet, partnerbijdragen, vrijwilligersinzet, bijdragen in natura en aangevraagde of beoogde ondersteuning via andere sport- en maatschappelijke fondsen. De bijdrage wordt ingezet binnen de jaarpilot in regio 's-Hertogenbosch.",
    )

    add_heading(doc, "13. Relevantie voor Rabobank")
    add_p(doc, "Het Velo Vital Leenfiets & Community Programma sluit sterk aan bij de maatschappelijke thema's van Rabobank:")
    add_bullets(doc, [
        "Gezond leven en sterke buurten: vrouwen worden gestimuleerd om meer te bewegen en elkaar lokaal te ontmoeten.",
        "Sociale verbinding: deelnemers stappen in binnen een toegankelijke community in plaats van alleen.",
        "Duurzame mobiliteit: fietsen wordt gestimuleerd als gezonde en duurzame vorm van verplaatsing en vrijetijdsbesteding.",
        "Structurele positieve verandering: het project ontwikkelt een herhaalbare pilotaanpak waarmee deelnemers na de eerste kennismaking blijven bewegen en waarmee Velo Vital later ook in andere regio's maatschappelijke impact kan realiseren.",
    ])
    add_p(doc, "Het project is geen commerciële fietsverhuur, maar een maatschappelijk programma dat vrouwen helpt om van interesse naar structureel beweeggedrag te komen.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
