from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Velo_Vital_ABN_AMRO_Fonds_Subsidieaanvraag.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 37, 52)
MUTED = RGBColor(90, 96, 108)
ORANGE = RGBColor(226, 105, 47)
LIGHT = "F4F6F9"
LIGHT_BLUE = "E8EEF5"
GRID = "D9DCE3"
WHITE = "FFFFFF"


def qstyle(style, font="Calibri", size=11, color=INK, bold=None):
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:ascii"), font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_tbl_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def set_column_widths(table, widths_dxa):
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def paragraph_border_bottom(paragraph, color="E2692F", size="10", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_p(doc, text="", after=6, before=0, line=1.25, color=INK, bold=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    if text:
        r = p.add_run(text)
        set_run(r, size=11, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run(run, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run(run, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run(r, size=10.7, color=INK)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_tbl_width(table, 9360, 120)
    set_table_borders(table, "E2E5EA", "4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, 160, 200, 160, 200)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + " ")
    set_run(r1, size=10.7, color=ORANGE, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=10.7, color=INK)
    add_p(doc, "", after=4)


def add_metadata_table(doc):
    rows = [
        ("Organisatie", "Stichting Wielerevenementen Aalburg"),
        ("Project", "Velo Vital Leenfiets & Community Programma"),
        ("Rechtsvorm", "Stichting"),
        ("KvK-nummer", "60618167"),
        ("Vestigingsadres", "St.-Lambertusstraat 65, 5266AD Cromvoirt"),
        ("Projectgebied", "'s-Hertogenbosch en omgeving"),
        ("Projectperiode", "September 2026 tot en met augustus 2027"),
        ("Gevraagde bijdrage", "EUR 15.000"),
        ("Totale projectwaarde", "EUR 30.000"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_tbl_width(table, 9360, 120)
    set_column_widths(table, [2600, 6760])
    set_table_borders(table)
    for label, value in rows:
        c1, c2 = table.rows[rows.index((label, value))].cells
        set_cell_shading(c1, LIGHT_BLUE)
        set_cell_margins(c1)
        set_cell_margins(c2)
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run(r1, size=10.2, color=DARK_BLUE, bold=True)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run(r2, size=10.2, color=INK, bold=label in ("Gevraagde bijdrage", "Totale projectwaarde"))
    add_p(doc, "", after=8)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_tbl_width(table, 9360, 120)
    set_column_widths(table, widths)
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run(r, size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i == len(row) - 1 and value.startswith("EUR"):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(value)
            set_run(r, size=9.7, color=INK, bold=value.startswith("Totale") or value.startswith("EUR 30.000"))
    add_p(doc, "", after=8)
    return table


def add_cover(doc):
    add_p(doc, "Stichting Wielerevenementen Aalburg", after=8, color=MUTED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Velo Vital Leenfiets & Community Programma")
    set_run(r, size=24, color=INK, bold=True)
    add_p(doc, "Subsidieaanvraag ABN AMRO Fonds", after=8, color=MUTED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Laagdrempelig sportief fietsen voor vrouwen in 's-Hertogenbosch en omgeving", after=18, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="E2692F", size="12", space="6")
    add_p(doc, "", after=10)
    add_metadata_table(doc)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    qstyle(styles["Normal"], size=11, color=INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    qstyle(styles["Heading 1"], size=16, color=BLUE, bold=True)
    qstyle(styles["Heading 2"], size=13, color=BLUE, bold=True)
    qstyle(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True)
    qstyle(styles["List Bullet"], size=10.7, color=INK)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.194)
    styles["List Bullet"].paragraph_format.space_after = Pt(4)
    styles["List Bullet"].paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.text = "Velo Vital | Subsidieaanvraag ABN AMRO Fonds"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "Stichting Wielerevenementen Aalburg"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.runs[0], size=9, color=MUTED)

    add_cover(doc)
    add_callout(
        doc,
        "Kernvraag:",
        "Een bijdrage van EUR 15.000 maakt het mogelijk om vrouwen zonder grote investering, zonder prestatiedruk en met begeleiding te laten starten met sportief fietsen.",
    )

    add_heading(doc, "1. Korte samenvatting")
    add_p(doc, "Velo Vital wil vrouwen in 's-Hertogenbosch en omgeving helpen om op een laagdrempelige manier te starten met sportief fietsen. Voor veel vrouwen is de eerste stap groot: een geschikte fiets is duur, technische kennis ontbreekt, er is geen fietsmaatje of groep, en de sportcultuur kan voor beginners prestatiegericht of weinig toegankelijk voelen.")
    add_p(doc, "Met het Velo Vital Leenfiets & Community Programma kunnen vrouwen een fiets tijdelijk lenen voor een periode van een tot drie maanden. Zij krijgen begeleiding, toegang tot laagdrempelige fietsmomenten en aansluiting bij een community waarin plezier, veiligheid, zelfvertrouwen en ontmoeting centraal staan.")

    add_heading(doc, "2. Maatschappelijke aanleiding")
    add_p(doc, "Sportief fietsen heeft veel maatschappelijke waarde. Het combineert bewegen, buiten zijn, gezondheid, duurzame mobiliteit en sociale ontmoeting. Toch is de toegang niet vanzelfsprekend. Anders dan bij veel andere sporten vraagt fietsen vaak direct om materiaal, kennis, vertrouwen en een netwerk.")
    add_bullets(doc, [
        "hoge opstartkosten voor fiets, helm, kleding en accessoires;",
        "onzekerheid over materiaal, onderhoud en veiligheid;",
        "geen fietsmaatje of laagdrempelige groep om mee te starten;",
        "twijfel of sportief fietsen past bij conditie, ervaring of leefritme;",
        "een sportcultuur die voor beginners technisch of prestatiegericht kan voelen;",
        "gebrek aan een herkenbare, vrouwgerichte instaproute.",
    ])
    add_p(doc, "Hierdoor blijft een groep vrouwen aan de kant staan, terwijl zij juist baat kan hebben bij meer beweging, sociale verbinding en groeiend zelfvertrouwen. Velo Vital maakt de eerste stap concreet: eerst proberen, ervaren en vertrouwen opbouwen, zonder directe aanschaf en zonder prestatiedruk.")

    add_heading(doc, "3. Doelstelling en doelgroep")
    add_p(doc, "Het doel van het project is om vrouwen in 's-Hertogenbosch en omgeving laagdrempelig te laten kennismaken met sportief fietsen en hen te begeleiden naar structureel beweeggedrag.")
    add_bullets(doc, [
        "gelijke toegang tot sportief fietsen voor vrouwen die door kosten, materiaal of onzekerheid niet starten;",
        "meer beweging, vitaliteit en mentale veerkracht;",
        "sociale verbinding via een herkenbare lokale community;",
        "zelfvertrouwen en zelfstandigheid in bewegen;",
        "doorstroom naar bestaand lokaal sport- en beweegaanbod.",
    ])
    add_p(doc, "De primaire doelgroep bestaat uit volwassen vrouwen in 's-Hertogenbosch en omgeving die interesse hebben in sportief fietsen, maar de stap nog niet zetten. Binnen de communicatie kan ook ruimte worden gemaakt voor jonge vrouwen vanaf 16 jaar die dreigen af te haken bij sport of moeite hebben om passend aanbod te vinden.")

    add_heading(doc, "4. Projectaanpak")
    for title, text in [
        ("Leenfietsprogramma", "Velo Vital stelt fietsen beschikbaar aan vrouwen die sportief fietsen willen proberen. Deelnemers kunnen een fiets tijdelijk gebruiken voor een periode van een tot drie maanden."),
        ("Intake en persoonlijke matching", "Iedere deelnemer krijgt een korte intake. Daarbij kijken we naar ervaring, motivatie, beschikbare tijd, eventuele onzekerheden, gewenste intensiteit en praktische behoeften."),
        ("Community en laagdrempelige fietsmomenten", "Deelnemers worden onderdeel van de Velo Vital-community. Zij kunnen aansluiten bij rustige fietsmomenten en contact leggen met andere vrouwen die ook willen starten of al fietsen."),
        ("Fietschecks, onderhoud en basismateriaal", "Velo Vital werkt met een fietsenmaker of materiaalpartner die fietsen beschikbaar stelt, onderhoud uitvoert en zorgt dat het materiaal veilig en gebruiksklaar blijft."),
        ("Lokale campagne en bereik", "De campagne bestaat uit lokale online zichtbaarheid, herkenbare Velo Vital-content, promotie via partners, deelnemersverhalen en laagdrempelige kennismakingsmomenten."),
        ("Doorstroom naar blijvend bewegen", "Aan het einde van de leenperiode bespreekt Velo Vital met iedere deelnemer wat een passende vervolgstap is, zoals zelfstandig blijven fietsen of aansluiten bij lokaal aanbod."),
    ]:
        add_heading(doc, title, level=3)
        add_p(doc, text)

    add_heading(doc, "5. Beoogde resultaten")
    add_bullets(doc, [
        "minimaal 25 vrouwen nemen deel aan een leenfietsperiode van een tot drie maanden;",
        "minimaal 75 vrouwen worden actief betrokken via community-activiteiten, kennismakingsmomenten, proefritten of campagne;",
        "minimaal 500 vrouwen worden bereikt via lokale communicatie, sociale media, partners en mond-tot-mondwerving;",
        "minimaal 20 van de 25 leenfietsdeelnemers blijven na afloop actief fietsen;",
        "minimaal 12 deelnemers stromen door naar structureel aanbod;",
        "minimaal 12 laagdrempelige community- of kennismakingsmomenten worden georganiseerd;",
        "er wordt een herhaalbare pilotaanpak ontwikkeld voor leenfietsen, intake, communitybinding, onderhoud en doorstroom.",
    ])

    add_heading(doc, "6. Aansluiting bij ABN AMRO Fonds")
    add_p(doc, "Velo Vital sluit aan bij een maatschappelijke fondsbenadering waarin gelijke kansen, lokale betrokkenheid en het versterken van mensen centraal staan. Het project verlaagt financiële en praktische drempels, maar richt zich ook op sociale drempels: niet weten waar je moet beginnen, niemand kennen om mee te gaan, of onzeker zijn over je eigen niveau.")
    add_p(doc, "Een bijdrage van ABN AMRO Fonds maakt het mogelijk om vrouwen niet alleen tijdelijk materiaal te bieden, maar ook begeleiding, community, zichtbaarheid en een route naar blijvend bewegen. Daarmee ontstaat maatschappelijke impact die verder gaat dan een losse activiteit.")

    add_heading(doc, "7. Planning")
    add_simple_table(doc, ["Periode", "Activiteiten"], [
        ("Juli - augustus 2026", "Voorbereiding, partnerafspraken, selectie fietsen, opzetten aanmeldproces, campagnevoorbereiding"),
        ("September - november 2026", "Start lokale campagne, werving deelnemers, intakegesprekken, eerste leenfietsronde en community-activiteiten"),
        ("December 2026 - februari 2027", "Tweede leenfietsronde, winterproof community-aanbod, begeleiding, onderhoud en tussentijdse evaluatie"),
        ("Maart - mei 2027", "Derde leenfietsronde, extra kennismakingsmomenten, doorstroom naar structureel aanbod en partneractivatie"),
        ("Juni - augustus 2027", "Afronding jaarpilot, impactmeting, financiële afronding, terugkoppeling aan financiers en uitwerking opschalingslessen"),
    ], [2500, 6860])

    add_heading(doc, "8. Samenwerking en partners")
    add_bullets(doc, [
        "fietsenmaker of materiaalpartner voor fietsen, onderhoud en technische expertise;",
        "Vrouwenwielrennen Den Bosch als mogelijke doorstroompartner;",
        "S-PORT en Bosch Sportakkoord voor lokale sportverbinding en zichtbaarheid;",
        "Den Bosch City voor lokale zichtbaarheid, netwerkbereik en promotionele ondersteuning;",
        "wijk-, welzijns- en preventiepartners voor het bereiken van vrouwen die niet vanzelf bij sportaanbod terechtkomen;",
        "vrijwilligers en communitybegeleiders voor laagdrempelige begeleiding en fietsmomenten.",
    ])

    add_heading(doc, "9. Monitoring en evaluatie")
    add_p(doc, "Velo Vital meet de voortgang praktisch en resultaatgericht. We registreren aanmeldingen, daadwerkelijke leenfietsdeelname, duur van de leenperiode, deelname aan community-activiteiten, doorstroom naar structureel aanbod en ervaringen van deelnemers via korte testimonials.")
    add_p(doc, "Succes betekent dat een deelnemer na de leenperiode blijft fietsen en zich zekerder voelt om sportief te bewegen, via Velo Vital, een andere fietscommunity of zelfstandig.")

    add_heading(doc, "10. Begroting")
    add_heading(doc, "Kosten", level=2)
    add_simple_table(doc, ["Kostenpost", "Bedrag"], [
        ("Lokale marketingcampagne en zichtbaarheid voor deelnemerswerving", "EUR 5.000"),
        ("Content, vormgeving en deelnemersverhalen voor social media en partners", "EUR 3.000"),
        ("Kennismakingsritten en community-activatie voor nieuwe deelnemers", "EUR 3.500"),
        ("Fietschecks, onderhoud en rijklaar maken van leenfietsen", "EUR 4.500"),
        ("Praktisch instapmateriaal voor deelnemers", "EUR 3.000"),
        ("Aanmeldproces, intake en praktische deelnemersondersteuning", "EUR 2.500"),
        ("Monitoring, evaluatie en korte impactrapportage", "EUR 1.000"),
        ("Inzet fietsenmaker of materiaalpartner in natura", "EUR 3.500"),
        ("Vrijwilligersinzet community en begeleiding in natura", "EUR 2.000"),
        ("Inzet lokale partners voor zichtbaarheid en promotionele ondersteuning in natura", "EUR 1.500"),
        ("Eigen inzet Velo Vital", "EUR 500"),
        ("Totale projectwaarde", "EUR 30.000"),
    ], [7060, 2300])

    add_heading(doc, "Dekking", level=2)
    add_simple_table(doc, ["Dekking", "Bedrag"], [
        ("Gevraagde bijdrage ABN AMRO Fonds", "EUR 15.000"),
        ("Eigen inzet Velo Vital", "EUR 500"),
        ("Bijdrage fietsenmaker of materiaalpartner in natura", "EUR 3.500"),
        ("Vrijwilligersinzet community en begeleiding", "EUR 2.000"),
        ("Bijdrage lokale partners in natura", "EUR 1.500"),
        ("Aangevraagde of beoogde bijdrage sport- en maatschappelijke fondsen", "EUR 7.500"),
        ("Totale dekking", "EUR 30.000"),
    ], [7060, 2300])

    add_heading(doc, "11. Gevraagde bijdrage en borging")
    add_p(doc, "Velo Vital vraagt een bijdrage van EUR 15.000 aan ABN AMRO Fonds. Deze bijdrage wordt ingezet voor zichtbare en direct maatschappelijke projectactiviteiten: deelnemerswerving, communicatie, content, kennismakingsritten, community-activatie, fietschecks, onderhoud, instapmateriaal en praktische deelnemersondersteuning.")
    add_p(doc, "De investering in leenfietsen, onderhoudsafspraken, intake, communitystructuur en partnernetwerk blijft na de projectperiode bruikbaar. De pilot levert een herhaalbaar model op waarmee Velo Vital ook na afloop vrouwen kan begeleiden van eerste kennismaking naar blijvend fietsen.")

    add_heading(doc, "12. Communicatie")
    add_p(doc, "Bij een bijdrage van ABN AMRO Fonds benoemt Velo Vital ABN AMRO Fonds als maatschappelijke partner in passende communicatie-uitingen. Velo Vital deelt voortgang, beeldmateriaal, deelnemersverhalen en resultaten met het fonds en draagt actief bij aan zichtbaarheid van de maatschappelijke impact.")
    add_callout(doc, "Kernboodschap:", "Velo Vital maakt sportief fietsen toegankelijk voor vrouwen die willen bewegen, maar nog niet weten waar of hoe ze moeten beginnen. Zonder grote investering, zonder prestatiedruk en met een community om op terug te vallen.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
